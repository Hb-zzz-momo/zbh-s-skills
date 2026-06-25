#!/usr/bin/env python3
"""Audit basic formatting signals in a Word .docx paper."""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:  # pragma: no cover
    print("Missing dependency: python-docx. Install with: python -m pip install python-docx", file=sys.stderr)
    raise SystemExit(2)


CAPTION_RE = re.compile(r"^\s*(图|表|Figure|Table)\s*[\d一二三四五六七八九十IVXivx\-\.]*")
REF_HEADING_RE = re.compile(r"^\s*(参考文献|References|Bibliography)\s*$", re.I)


def emu_to_cm(value: Any) -> float | None:
    if value is None:
        return None
    return round(float(value) / 360000.0, 2)


def pt_value(value: Any) -> float | None:
    if value is None:
        return None
    return round(float(value.pt), 2)


def alignment_name(value: Any) -> str | None:
    if value is None:
        return None
    return str(value).split(".")[-1]


def line_spacing_value(value: Any) -> float | str | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return value
    if hasattr(value, "pt"):
        return f"{round(float(value.pt), 2)} pt"
    return str(value)


def paragraph_text(paragraph: Any) -> str:
    return paragraph.text.replace("\u3000", " ").strip()


def collect_run_sizes(doc: Any) -> Counter:
    sizes: Counter = Counter()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() and run.font.size is not None:
                sizes[str(pt_value(run.font.size))] += 1
    return sizes


def audit_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraph_texts = [paragraph_text(p) for p in doc.paragraphs]
    blank_runs = []
    current_blank_start = None
    current_blank_count = 0
    for idx, text in enumerate(paragraph_texts, start=1):
        if not text:
            if current_blank_start is None:
                current_blank_start = idx
            current_blank_count += 1
        else:
            if current_blank_count > 1:
                blank_runs.append({"start": current_blank_start, "count": current_blank_count})
            current_blank_start = None
            current_blank_count = 0
    if current_blank_count > 1:
        blank_runs.append({"start": current_blank_start, "count": current_blank_count})

    sections = []
    for idx, section in enumerate(doc.sections, start=1):
        sections.append(
            {
                "index": idx,
                "start_type": str(section.start_type).split(".")[-1],
                "page_width_cm": emu_to_cm(section.page_width),
                "page_height_cm": emu_to_cm(section.page_height),
                "top_margin_cm": emu_to_cm(section.top_margin),
                "bottom_margin_cm": emu_to_cm(section.bottom_margin),
                "left_margin_cm": emu_to_cm(section.left_margin),
                "right_margin_cm": emu_to_cm(section.right_margin),
                "header_distance_cm": emu_to_cm(section.header_distance),
                "footer_distance_cm": emu_to_cm(section.footer_distance),
            }
        )

    style_names = ["Normal", "Heading 1", "Heading 2", "Heading 3"]
    styles = {}
    for name in style_names:
        try:
            style = doc.styles[name]
        except KeyError:
            styles[name] = {"present": False}
            continue
        fmt = getattr(style, "paragraph_format", None)
        styles[name] = {
            "present": True,
            "font_name": style.font.name,
            "font_size_pt": pt_value(style.font.size),
            "bold": style.font.bold,
            "italic": style.font.italic,
            "line_spacing": line_spacing_value(getattr(fmt, "line_spacing", None)),
            "space_before_pt": pt_value(getattr(fmt, "space_before", None)),
            "space_after_pt": pt_value(getattr(fmt, "space_after", None)),
            "first_line_indent_cm": emu_to_cm(getattr(fmt, "first_line_indent", None)),
            "alignment": alignment_name(getattr(fmt, "alignment", None)),
        }

    captions = []
    reference_heading_indexes = []
    for idx, text in enumerate(paragraph_texts, start=1):
        if CAPTION_RE.match(text):
            captions.append({"paragraph": idx, "text": text[:120]})
        if REF_HEADING_RE.match(text):
            reference_heading_indexes.append(idx)

    tables = []
    for idx, table in enumerate(doc.tables, start=1):
        row_count = len(table.rows)
        col_count = len(table.columns)
        cell_text_lengths = [len(cell.text.strip()) for row in table.rows for cell in row.cells]
        tables.append(
            {
                "index": idx,
                "rows": row_count,
                "columns": col_count,
                "style": table.style.name if table.style is not None else None,
                "empty_cells": sum(1 for length in cell_text_lengths if length == 0),
            }
        )

    issues = []
    if blank_runs:
        issues.append({"severity": "warning", "message": "Found consecutive blank paragraphs.", "details": blank_runs[:20]})
    if not reference_heading_indexes:
        issues.append({"severity": "info", "message": "No obvious References heading found."})
    if not captions:
        issues.append({"severity": "info", "message": "No obvious figure/table captions found."})

    return {
        "file": str(path.resolve()),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "sections": sections,
        "styles": styles,
        "run_font_sizes": dict(collect_run_sizes(doc).most_common()),
        "consecutive_blank_paragraphs": blank_runs,
        "captions": captions[:50],
        "reference_headings": reference_heading_indexes,
        "tables": tables,
        "issues": issues,
    }


def to_markdown(report: dict[str, Any]) -> str:
    lines = [
        "# DOCX 格式审计报告",
        "",
        f"- 文件：`{report['file']}`",
        f"- 段落数：{report['paragraph_count']}",
        f"- 表格数：{report['table_count']}",
        "",
        "## 页面设置",
        "",
    ]
    for section in report["sections"]:
        lines.append(
            "- Section {index}: {page_width_cm}x{page_height_cm} cm, margins "
            "T/B/L/R={top_margin_cm}/{bottom_margin_cm}/{left_margin_cm}/{right_margin_cm} cm, "
            "header/footer={header_distance_cm}/{footer_distance_cm} cm".format(**section)
        )
    lines.extend(["", "## 样式", ""])
    for name, data in report["styles"].items():
        if not data.get("present"):
            lines.append(f"- `{name}`: missing")
            continue
        lines.append(
            f"- `{name}`: font={data['font_name']}, size={data['font_size_pt']} pt, "
            f"line_spacing={data['line_spacing']}, before/after={data['space_before_pt']}/{data['space_after_pt']} pt, "
            f"first_indent={data['first_line_indent_cm']} cm, align={data['alignment']}"
        )
    lines.extend(["", "## 字号分布", ""])
    if report["run_font_sizes"]:
        for size, count in report["run_font_sizes"].items():
            lines.append(f"- {size} pt: {count} runs")
    else:
        lines.append("- No direct run font sizes found.")
    lines.extend(["", "## 图表题注线索", ""])
    for item in report["captions"] or []:
        lines.append(f"- P{item['paragraph']}: {item['text']}")
    if not report["captions"]:
        lines.append("- 未发现明显图表题注。")
    lines.extend(["", "## 参考文献线索", ""])
    if report["reference_headings"]:
        lines.append("- Heading paragraphs: " + ", ".join(map(str, report["reference_headings"])))
    else:
        lines.append("- 未发现明显参考文献标题。")
    lines.extend(["", "## 表格", ""])
    for table in report["tables"]:
        lines.append(f"- Table {table['index']}: {table['rows']}x{table['columns']}, style={table['style']}, empty_cells={table['empty_cells']}")
    if not report["tables"]:
        lines.append("- No tables found.")
    lines.extend(["", "## 问题", ""])
    for issue in report["issues"]:
        lines.append(f"- [{issue['severity']}] {issue['message']}")
    if not report["issues"]:
        lines.append("- 未发现脚本可判定的问题。")
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Audit formatting signals in a .docx paper.")
    parser.add_argument("--input", required=True, help="Input .docx path")
    parser.add_argument("--out", required=True, help="Output .md or .json report path")
    args = parser.parse_args()

    input_path = Path(args.input)
    out_path = Path(args.out)
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2
    if input_path.suffix.lower() != ".docx":
        print("Input must be a .docx file.", file=sys.stderr)
        return 2

    report = audit_docx(input_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if out_path.suffix.lower() == ".json":
        out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    else:
        out_path.write_text(to_markdown(report), encoding="utf-8")
    print(str(out_path.resolve()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
