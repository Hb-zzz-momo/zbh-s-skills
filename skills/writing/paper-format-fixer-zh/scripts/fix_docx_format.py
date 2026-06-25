#!/usr/bin/env python3
"""Conservatively fix formatting in a Word .docx paper."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:  # pragma: no cover
    print("Missing dependency: python-docx. Install with: python -m pip install python-docx", file=sys.stderr)
    raise SystemExit(2)


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def load_rules(path: Path | None) -> dict[str, Any]:
    if path is None:
        return {}
    if not path.exists():
        raise FileNotFoundError(f"Rules file not found: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_can_write(path: Path, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise FileExistsError(f"Output already exists: {path}. Use --overwrite to replace it.")
    path.parent.mkdir(parents=True, exist_ok=True)


def set_east_asia_font(element: Any, font_name: str) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def apply_style_rule(style: Any, rule: dict[str, Any]) -> None:
    if "font_name" in rule:
        style.font.name = rule["font_name"]
    if "east_asia_font" in rule:
        set_east_asia_font(style._element, rule["east_asia_font"])
    if "font_size_pt" in rule:
        style.font.size = Pt(float(rule["font_size_pt"]))
    if "bold" in rule:
        style.font.bold = bool(rule["bold"])
    if "italic" in rule:
        style.font.italic = bool(rule["italic"])

    fmt = style.paragraph_format
    if "line_spacing" in rule:
        fmt.line_spacing = float(rule["line_spacing"])
    if "space_before_pt" in rule:
        fmt.space_before = Pt(float(rule["space_before_pt"]))
    if "space_after_pt" in rule:
        fmt.space_after = Pt(float(rule["space_after_pt"]))
    if "first_line_indent_cm" in rule:
        fmt.first_line_indent = Cm(float(rule["first_line_indent_cm"]))
    if "left_indent_cm" in rule:
        fmt.left_indent = Cm(float(rule["left_indent_cm"]))
    if "right_indent_cm" in rule:
        fmt.right_indent = Cm(float(rule["right_indent_cm"]))
    if "alignment" in rule:
        alignment = ALIGNMENTS.get(str(rule["alignment"]).lower())
        if alignment is not None:
            fmt.alignment = alignment


def apply_page_rules(doc: Any, page_rule: dict[str, Any]) -> int:
    changed = 0
    mapping = {
        "top_cm": "top_margin",
        "bottom_cm": "bottom_margin",
        "left_cm": "left_margin",
        "right_cm": "right_margin",
        "header_cm": "header_distance",
        "footer_cm": "footer_distance",
    }
    for section in doc.sections:
        for key, attr in mapping.items():
            if key in page_rule:
                setattr(section, attr, Cm(float(page_rule[key])))
                changed += 1
    return changed


def collapse_blank_paragraphs(doc: Any) -> int:
    removed = 0
    previous_blank = False
    for paragraph in list(doc.paragraphs):
        is_blank = not paragraph.text.strip()
        if is_blank and previous_blank:
            element = paragraph._element
            element.getparent().remove(element)
            removed += 1
            continue
        previous_blank = is_blank
    return removed


def normalize_abnormal_run_sizes(doc: Any, rule: dict[str, Any]) -> int:
    min_pt = float(rule.get("min_pt", 6))
    max_pt = float(rule.get("max_pt", 36))
    replacement = Pt(float(rule.get("replacement_pt", 12)))
    changed = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.text.strip() or run.font.size is None:
                continue
            size = run.font.size.pt
            if size < min_pt or size > max_pt:
                run.font.size = replacement
                changed += 1
    return changed


def apply_table_style(doc: Any, style_name: str) -> int:
    changed = 0
    for table in doc.tables:
        try:
            table.style = style_name
            changed += 1
        except KeyError:
            continue
    return changed


def fix_docx(input_path: Path, out_path: Path, rules: dict[str, Any], overwrite: bool) -> dict[str, Any]:
    ensure_can_write(out_path, overwrite)
    doc = Document(str(input_path))
    summary = {
        "input": str(input_path.resolve()),
        "output": str(out_path.resolve()),
        "page_properties_changed": 0,
        "styles_changed": [],
        "blank_paragraphs_removed": 0,
        "abnormal_run_sizes_changed": 0,
        "tables_styled": 0,
        "mode": "rules" if rules else "conservative-default",
    }

    if "page" in rules:
        summary["page_properties_changed"] = apply_page_rules(doc, rules["page"])

    for style_name, rule in rules.get("styles", {}).items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        apply_style_rule(style, rule)
        summary["styles_changed"].append(style_name)

    if rules.get("collapse_blank_paragraphs", True):
        summary["blank_paragraphs_removed"] = collapse_blank_paragraphs(doc)

    run_font_rule = rules.get("run_font", {"min_pt": 6, "max_pt": 36, "replacement_pt": 12})
    summary["abnormal_run_sizes_changed"] = normalize_abnormal_run_sizes(doc, run_font_rule)

    if "table_style" in rules:
        summary["tables_styled"] = apply_table_style(doc, str(rules["table_style"]))

    doc.save(str(out_path))
    Document(str(out_path))
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description="Conservatively fix formatting in a .docx paper.")
    parser.add_argument("--input", required=True, help="Input .docx path")
    parser.add_argument("--rules", help="Optional JSON rules path")
    parser.add_argument("--out", required=True, help="Output .docx path")
    parser.add_argument("--overwrite", action="store_true", help="Allow replacing an existing output file")
    args = parser.parse_args()

    input_path = Path(args.input)
    out_path = Path(args.out)
    rules_path = Path(args.rules) if args.rules else None
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2
    if input_path.suffix.lower() != ".docx" or out_path.suffix.lower() != ".docx":
        print("Input and output must be .docx files.", file=sys.stderr)
        return 2

    try:
        rules = load_rules(rules_path)
        summary = fix_docx(input_path, out_path, rules, args.overwrite)
    except Exception as exc:
        print(f"Failed: {exc}", file=sys.stderr)
        return 1
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
